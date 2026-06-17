from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Estilos base ────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, text, level, color_hex):
    run = paragraph.add_run(text)
    run.bold = True
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)

def h1(text, color='1C1F3A'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    set_heading(p, text, 1, color)
    return p

def h2(text, color='CD1719'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    set_heading(p, text, 2, color)
    return p

def h3(text, color='575756'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    set_heading(p, text, 3, color)
    return p

def h4(text, color='9D9D9C'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_heading(p, text, 4, color)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def route_tag(route):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'Ruta: {route}')
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x11, 0x55, 0xCC)
    return p

def divider():
    p = doc.add_paragraph('─' * 80)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('CISP Colombia')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x9D, 0x9D, 0x9C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('Documentación de Sitios Web\nObservatorios CISP')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1C, 0x1F, 0x3A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('Estructura, secciones y descripción de cada sitio')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x57, 0x57, 0x56)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x9D, 0x9D, 0x9C)

page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SITIO 1 — TEJIENDO TERRITORIOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('SITIO 1 — Observatorio Tejiendo Territorios', '1C1F3A')
body('Plataforma de seguimiento y monitoreo de iniciativas comunitarias indígenas en Colombia. Documenta trayectorias organizativas, territorio, memoria y herramientas de evaluación para comunidades y organizaciones participantes del convenio.')
divider()

h2('Navegación principal', 'FF9331')
body('El sitio cuenta con una barra de navegación fija que enlaza a las secciones principales: Inicio, Sobre el Convenio, Tejiendo Trayectorias, Iniciativas Comunitarias e Índice de Cultura y Educación.')

# — Página Inicio
h2('1. Inicio (Home)', '1C1F3A')
route_tag('/')
body('Página de bienvenida del observatorio. Presenta el propósito del convenio y ofrece acceso rápido a las secciones principales.')
h3('Secciones de la página:')
bullet('Carrusel de imágenes — Galería fotográfica rotativa que ilustra las comunidades e iniciativas del convenio.')
bullet('Tarjetas de acceso rápido — Six cards que enlazan a: Sobre el Convenio, Tejiendo Trayectorias, Iniciativas Comunitarias, Índice de Cultura y Educación, Visualización y Consulta, y Reportes y Productos.')
bullet('Pie de página — Información del proyecto e instituciones aliadas.')

# — Sobre el Convenio
h2('2. Sobre el Convenio', '1C1F3A')
route_tag('/sobre-el-convenio')
body('Sección informativa sobre el convenio interinstitucional. Describe los objetivos, cobertura territorial y actores involucrados.')
h3('Secciones de la página:')
bullet('Descripción general — Texto con objetivos y alcance del convenio.')
bullet('Cobertura — Información sobre territorios y departamentos cubiertos.')
bullet('Actores — Listado de instituciones y organizaciones participantes.')
bullet('Documentos descargables — Cards con PDFs organizados por categoría: Convenio, Planificación, Diagnóstico, Normativo, Metodología. Cada card incluye título, descripción en modal y enlace de descarga.')
bullet('Video institucional — Sección con reproductor o placeholder de video.')

# — Tejiendo Trayectorias
h2('3. Tejiendo Trayectorias', '1C1F3A')
route_tag('/tejiendo-trayectorias')
body('Sección de seguimiento y monitoreo de las trayectorias organizativas. Permite explorar estadísticas y avances de las iniciativas agrupadas por líneas estratégicas.')
h3('Secciones de la página:')
bullet('Buscador y filtros — Búsqueda por texto, línea estratégica, tipo de iniciativa y región.')
bullet('Galería de tarjetas — Cards por iniciativa con fotografías y datos clave.')
bullet('Dashboard de estadísticas — Gráficas y métricas de avance por línea estratégica.')
bullet('Líneas estratégicas (5) con sus iniciativas:')
bullet('Economías Propias', 1)
bullet('Territorio y Medio Ambiente', 1)
bullet('Cultura e Identidad', 1)
bullet('Gobernanza y Gobierno Propio', 1)
bullet('Salud y Bienestar', 1)
bullet('Tipos de filtro — Acuerdo y Fortalecimiento.')

# — Iniciativas Comunitarias
h2('4. Iniciativas Comunitarias', '1C1F3A')
route_tag('/iniciativas-comunitarias')
body('Catálogo interactivo de todas las iniciativas comunitarias activas. Permite buscar, filtrar y explorar en mapa.')
h3('Secciones de la página:')
bullet('Mapa interactivo de Colombia — Mapa con calor (heat map) que muestra la concentración de iniciativas por departamento. Al hacer clic muestra estadísticas departamentales.')
bullet('Filtros — Por texto, línea estratégica, tipo (Acuerdo/Fortalecimiento) y región (Andina, Caribe, Pacífica, Amazónica, Orinoquía).')
bullet('Galería de cards — Una card por iniciativa con nombre, organización, comunidad y fotografías.')
bullet('Leyenda del mapa — Escala de colores según concentración de iniciativas.')

# — Detalle de Iniciativa
h2('5. Detalle de Iniciativa', '1C1F3A')
route_tag('/iniciativas-comunitarias/iniciativa/:id')
body('Página hub de cada iniciativa individual. Desde aquí se navega a los cinco sub-módulos de profundización.')
h3('Contenido:')
bullet('Encabezado de la iniciativa — Nombre, organización, comunidad y fotografía representativa.')
bullet('Cards de los 5 sub-módulos — Acceso a cada módulo temático con descripción breve.')

divider()
h3('Sub-módulo 5.1 — Tejiendo mi Territorio')
route_tag('/iniciativas-comunitarias/iniciativa/:id/tejiendo-mi-territorio')
body('Perfil geográfico y territorial de la iniciativa.')
bullet('Mapa del departamento — Mapa estático con zoom personalizado según el departamento de la iniciativa.')
bullet('Galería fotográfica — Imágenes del territorio.')
bullet('Documentos descargables — PDFs con perfil territorial, información geográfica y otros documentos específicos de la iniciativa.')

h3('Sub-módulo 5.2 — Tejiendo mis Saberes')
route_tag('/iniciativas-comunitarias/iniciativa/:id/tejiendo-mis-saberes')
body('Repositorio de conocimientos y patrimonio cultural de la comunidad.')
bullet('Biblioteca temática — Recursos organizados en 6 categorías: Economías Propias, Cuidado y Autocuidado, Identidad Cultural y Buen Vivir, Territorio, Salud Integral, Justicia Propia.')
bullet('Videos y materiales — Contenido multimedia por categoría temática.')
bullet('Tarjetas de recursos — Descripción, íconos por categoría y documentos asociados.')

h3('Sub-módulo 5.3 — Tejiendo mi Memoria')
route_tag('/iniciativas-comunitarias/iniciativa/:id/tejiendo-mi-memoria')
body('Archivo fotográfico y documental de la iniciativa.')
bullet('Galería de fotografías — Registro visual del proceso comunitario.')
bullet('Documentos de memoria — PDFs descargables sobre historia y memoria de la iniciativa.')
bullet('Organización por secciones — Archivo fotográfico, documentación de memoria y registros históricos.')

h3('Sub-módulo 5.4 — Entretejiendo Caminos')
route_tag('/iniciativas-comunitarias/iniciativa/:id/entretejiendo-caminos')
body('Matriz de evaluación y seguimiento de resultados de la iniciativa.')
bullet('Tabla de seguimiento — Columnas: Componente, Indicador, Meta, Resultado, % de cumplimiento, Estado, Observaciones.')
bullet('Estados con código de color — Cumplido, En proceso, Por iniciar, En riesgo.')
bullet('Barras de progreso — Visualización del avance por componente.')
bullet('Documentos de evaluación — PDFs descargables con informes de seguimiento.')

h3('Sub-módulo 5.5 — Tejidos que Transforman')
route_tag('/iniciativas-comunitarias/iniciativa/:id/tejidos-que-transforman')
body('Herramientas, hitos y logros transformadores de la iniciativa.')
bullet('Inventario de herramientas — Cards por herramienta: Seguimiento, Registro, Archivo, Evaluación. Incluye título, ícono, descripción y estado de uso.')
bullet('Sección de hitos — Logros clave y momentos de transformación con línea de tiempo.')
bullet('Recursos descargables — PDFs y enlaces de contenido.')

# — Visualización y Consulta
h2('6. Visualización y Consulta', '1C1F3A')
route_tag('/visualizacion')
body('Módulo de exploración de datos (en desarrollo). Framework preparado para mapa interactivo y dashboard de datos territoriales.')
bullet('Placeholder de mapa y dashboard.')
bullet('Estado: estructura lista, contenido pendiente.')

# — Reportes y Productos
h2('7. Reportes y Productos', '1C1F3A')
route_tag('/reportes')
body('Repositorio de documentos y entregables del proyecto (en desarrollo).')
bullet('Placeholder para informes, cartillas y productos del proyecto.')
bullet('Estado: estructura lista, contenido pendiente.')

# — Índice Cultura y Educación
h2('8. Índice de Cultura y Educación', '1C1F3A')
route_tag('/indice-cultura-educacion')
body('Indicadores culturales y educativos para comunidades indígenas (en desarrollo).')
bullet('Métricas de vitalidad cultural e indicadores educativos.')
bullet('Estado: próximamente.')

# — Admin
h2('9. Panel Administración', '1C1F3A')
route_tag('/admin-iniciativas')
body('Interfaz administrativa protegida con login para gestionar los datos de las iniciativas.')
bullet('Login con usuario y contraseña — Roles: admin, editor, viewer.')
bullet('Tabla de iniciativas — Búsqueda, filtros, edición, visibilidad y estado.')
bullet('Gestión de sesión — Login/Logout con persistencia en sessionStorage.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SITIO 2 — OBSERVATORIO PLANES DE VIDA
# ═══════════════════════════════════════════════════════════════════════════════
h1('SITIO 2 — Observatorio de Planes de Vida', '1C1F3A')
body('Plataforma de gestión, visualización y documentación de Planes de Vida de comunidades indígenas en Colombia. Organiza la información por organizaciones y planes con cinco módulos temáticos consistentes.')
divider()

h2('Navegación principal', '3E7F00')
body('Barra lateral fija con acceso a: Inicio, Visualización, Reportes. Diseño con paleta verde-naranja por región (Andina, Caribe, Pacífica, Amazónica, Orinoquía).')

# — Inicio
h2('1. Inicio (Home)', '1C1F3A')
route_tag('/')
body('Página de bienvenida que contextualiza el observatorio y los Planes de Vida indígenas.')
h3('Secciones de la página:')
bullet('Hero / Contextualización — Declaración de misión con fondo degradado. Texto institucional sobre el observatorio.')
bullet('Definición de Planes de Vida — Contenido educativo explicando qué son y su importancia para los pueblos indígenas.')
bullet('Carrusel de imágenes — Galería fotográfica rotativa de comunidades.')
bullet('Tarjetas de acceso rápido — Enlace a Visualización y a Reportes.')
bullet('Sección de enlaces de interés — Links categorizados por tipo: Estadísticas, Institucional, Normativo, Organizaciones, Cooperación, Internacional.')

# — Visualización y Consulta
h2('2. Visualización y Consulta', '1C1F3A')
route_tag('/visualizacion')
body('Dashboard interactivo público para explorar organizaciones y planes de vida.')
h3('Secciones de la página:')
bullet('Buscador y filtros — Búsqueda por texto, pueblo indígena y estado del plan.')
bullet('Mapa interactivo de Colombia — Muestra datos departamentales con clic-a-diálogo: número de planes, comunidades, pueblos y estado. Color-coded por región.')
bullet('Grilla de cards de organizaciones — Headers con degradado de color, ícono, nombre del pueblo (etnía), región, número de planes y estado.')

# — Mi Territorio
h2('3. Mi Territorio', '1C1F3A')
route_tag('/mi-territorio')
body('Dashboard de información territorial general. Muestra contexto geográfico, poblacional e identidad cultural.')
h3('Secciones de la página:')
bullet('Mapa interactivo — Visualización georreferenciada del territorio.')
bullet('Contexto territorial — Cards con: departamento, municipios, área, región geográfica, territorios protegidos, ecosistemas.')
bullet('Datos demográficos — Población total, distribución por género, familias, cabezas de hogar, comunidades internas.')
bullet('Elementos de identidad — Pueblo indígena, lengua nativa, autoridad tradicional, marco normativo, prácticas rituales, sitios sagrados.')

# — Mi Plan de Vida (directorio)
h2('4. Mi Plan de Vida (Directorio)', '1C1F3A')
route_tag('/mi-plan-de-vida')
body('Directorio buscable de todos los Planes de Vida registrados en el observatorio.')
h3('Secciones de la página:')
bullet('Búsqueda y filtros — Por texto, estado (Activo, En formulación, En actualización, En revisión) y región geográfica.')
bullet('Grilla de cards de planes — Color-coded por región, con: badge de estado, íconos de población, patrones decorativos de fondo y nombre del plan.')

# — Detalle de Plan
h2('5. Detalle de Plan', '1C1F3A')
route_tag('/mi-plan-de-vida/plan/:id')
body('Hub de navegación para un plan específico. Muestra el perfil del plan y acceso a sus cinco módulos temáticos.')
h3('Contenido:')
bullet('Identificación del plan — ID, región, esquema de colores propio.')
bullet('Cards de los 5 módulos — Con nombre y descripción de cada módulo.')

divider()
h3('Módulo 5.1 — Mi Territorio (nivel plan)')
route_tag('/mi-plan-de-vida/plan/:id/mi-territorio')
body('Contexto territorial específico del plan.')
bullet('Narrativa territorial — Descripción geográfica y cultural del territorio.')
bullet('Información étnica y cultural — Datos de identidad del pueblo.')
bullet('Mapa estático del territorio.')
bullet('Documentos descargables — Caracterizaciones y PDFs territoriales.')

h3('Módulo 5.2 — Mi Plan de Vida (nivel plan)')
route_tag('/mi-plan-de-vida/plan/:id/mi-plan-de-vida-modulo')
body('Documentación completa del Plan de Vida.')
bullet('Estructura del plan — Visión general y líneas de intervención.')
bullet('Líneas de intervención — Documentación detallada por línea.')
bullet('Propuesta de formulación — Secciones de la propuesta participativa.')
bullet('Aplicaciones de innovación — Herramientas digitales del plan.')
bullet('Documentos finales — PDFs del Plan de Vida para descarga.')
bullet('Versiones históricas — Documentos de versiones anteriores.')

h3('Módulo 5.3 — Fortalecimiento (nivel plan)')
route_tag('/mi-plan-de-vida/plan/:id/fortalecimiento')
body('Recursos pedagógicos y metodológicos organizados en 5 ejes temáticos.')
bullet('Economías Propias')
bullet('Cuidado y Autocuidado')
bullet('Identidad Cultural y Buen Vivir Comunitario')
bullet('Igualdad de Género')
bullet('Derechos')
bullet('Cada eje incluye presentaciones, PDFs y videos de capacitación.')

h3('Módulo 5.4 — Tejiendo Caminos (nivel plan)')
route_tag('/mi-plan-de-vida/plan/:id/tejiendo-caminos')
body('Diagnóstico organizacional y rutas de mejoramiento.')
bullet('Resultados del diagnóstico — Estado actual de capacidades organizativas.')
bullet('Análisis organizacional — Gráficas y visualizaciones de capacidades.')
bullet('Rutas de mejoramiento — Recomendaciones y plan de acción.')
bullet('Reportes PDF — Informes de diagnóstico descargables.')

h3('Módulo 5.5 — Archivo Vivo (nivel plan)')
route_tag('/mi-plan-de-vida/plan/:id/archivo-vivo')
body('Archivo de memoria territorial y comunitaria.')
bullet('Videos de narrativa comunitaria — Lista con metadatos (duración, categoría, tipo) y reproductor.')
bullet('Materiales documentales — Colección de documentos clasificados.')
bullet('Sistema de clasificación — Organización por categorías temáticas.')
bullet('Prácticas de preservación — Herramientas de gestión y conservación del archivo.')

# — Detalle de Organización
h2('6. Detalle de Organización', '1C1F3A')
route_tag('/org/:slug')
body('Hub de una organización específica. Mismos cinco módulos que el Detalle de Plan pero a nivel organizacional.')
h3('Módulos disponibles (misma estructura que el Detalle de Plan):')
bullet('Mi Territorio — Perfil territorial de la organización.')
bullet('Mi Plan de Vida — Documentación del plan de la organización.')
bullet('Fortalecimiento — Recursos por eje temático para la organización.')
bullet('Tejiendo Caminos — Diagnóstico organizacional específico.')
bullet('Archivo Vivo — Memoria y archivo documental de la organización.')

# — Secciones en desarrollo
h2('7. Secciones en Desarrollo', '1C1F3A')
h3('Fortalecimiento (general)')
route_tag('/fortalecimiento')
body('Repositorio general de recursos pedagógicos y metodológicos. Estructura lista, contenido en desarrollo.')
h3('Tejiendo Caminos (general)')
route_tag('/tejiendo-caminos')
body('Herramienta general de diagnóstico organizacional. Estructura lista, contenido en desarrollo.')
h3('Archivo Vivo (general)')
route_tag('/archivo-vivo')
body('Espacio dinámico de memoria territorial. Estructura lista, contenido en desarrollo.')
h3('Reportes y Productos')
route_tag('/reportes')
body('Repositorio de informes técnicos, boletines, fichas territoriales y análisis temáticos. La mayoría de ítems marcados como "próximamente".')

# — Admin
h2('8. Administración', '1C1F3A')
route_tag('/admin')
body('Panel administrativo protegido con login para gestionar organizaciones, planes y configuración del sistema.')
bullet('Login — Autenticación con lista de usuarios predefinidos.')
bullet('Gestión de organizaciones — Tablas CRUD con operaciones de edición.')
bullet('Configuración — Iconos, regiones, estados y ajustes generales.')
bullet('Gestión de sesión — Logout con persistencia de sesión.')

page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  SITIO 3 — OBSERVATORIO DE ACCESO A LA REPARACIÓN (VICTIMAS)
# ═══════════════════════════════════════════════════════════════════════════════
h1('SITIO 3 — Observatorio de Acceso a la Reparación', '1C1F3A')
body('Plataforma de análisis territorial sobre el acceso a la indemnización administrativa de víctimas del conflicto armado en Colombia. Identifica barreras, trayectorias y dinámicas operativas institucionales relacionadas con el derecho a la reparación.')
divider()

h2('Navegación principal', 'CD1719')
body('Barra de navegación fija oscura con: Inicio, Reparación, Módulos, Glosario. Paleta institucional: rojo (#CD1719), gris oscuro (#575756), amarillo (#FEC707).')

# — Inicio
h2('1. Inicio (Home)', '1C1F3A')
route_tag('/')
body('Página de bienvenida que presenta el observatorio, sus objetivos y los hechos victimizantes reconocidos por la Ley 1448.')
h3('Secciones de la página:')
bullet('Hero — Descripción de la plataforma y su misión de análisis.')
bullet('Objetivo general — Texto sobre el análisis del acceso a la indemnización administrativa.')
bullet('Objetivos específicos — Cuatro objetivos con íconos: caracterizar dinámicas de acceso, analizar barreras y brechas, comprender trayectorias, construir capacidad institucional.')
bullet('Hechos victimizantes — Cards clicables por hecho (8 tipos). Al hacer clic se abre un diálogo con información del hecho, marco normativo y monto de indemnización.')

# — Comprendiendo la Reparación
h2('2. Comprendiendo la Reparación', '1C1F3A')
route_tag('/comprendiendo-reparacion')
body('Sección educativa sobre los derechos de las víctimas y el proceso de reparación.')
h3('Secciones de la página:')
bullet('Cuatro derechos de las víctimas — Verdad, Justicia, Reparación integral, Garantías de No Repetición.')
bullet('Etapas del proceso de reparación — Con descripción de cada etapa:')
bullet('Registro en RUV (Registro Único de Víctimas)', 1)
bullet('Radicación administrativa', 1)
bullet('Requisitos documentales', 1)
bullet('Valoración y priorización', 1)
bullet('Desembolso de la indemnización', 1)
bullet('Cada etapa incluye descripción, documentos requeridos e información de tiempos. Formato de acordeón expandible.')

# — Glosario
h2('3. Glosario', '1C1F3A')
route_tag('/glosario')
body('Glosario buscable de términos clave del proceso de reparación.')
h3('Secciones de la página:')
bullet('12 términos definidos — RUV, UARIV, Hecho victimizante, Indemnización, Ley 1448 de 2011, DIH, Reparación integral, Contactabilidad, Barrera de acceso, SNARIV, Decreto 4800 de 2011, Fondo para la Reparación.')
bullet('Buscador — Filtro por texto en tiempo real.')
bullet('Enlaces a instituciones externas — UARIV, Centro Nacional de Memoria Histórica, CISP América Latina.')

# — Módulos (hub)
h2('4. Módulos (Hub)', '1C1F3A')
route_tag('/modulos')
body('Página de entrada a los cinco módulos analíticos del observatorio.')
h3('Contenido:')
bullet('Cards de los 5 módulos — Numerados 01–05, con título, descripción, íconos y etiquetas temáticas.')
bullet('Interacción — Hover con animación, clic navega al módulo correspondiente.')

divider()

# — Módulo 01
h2('Módulo 01 — Radiografía del Acceso', '1C1F3A')
route_tag('/modulos/radiografia-acceso')
body('Análisis cuantitativo general del acceso a la indemnización por departamento.')
h3('Secciones del módulo:')
bullet('Mapa interactivo de Colombia — Heat map con intensidad por concentración de víctimas (< 80K, 80–200K, 200–400K, 400K+). Clic en departamento muestra panel de detalle.')
bullet('Panel de detalle departamental — Al seleccionar un departamento: total de víctimas en RUV, quiénes reclamaron vs. no reclamaron (con barra de progreso), distribución de hechos victimizantes con barras de frecuencia.')
bullet('Indicadores nacionales — Totales consolidados: víctimas en RUV, reclamaron, no reclamaron.')
bullet('Top 8 departamentos — Ranking por número de víctimas.')
bullet('Hechos victimizantes (8 tipos) — Desplazamiento forzado, Homicidio, Desaparición forzada, Secuestro, Reclutamiento ilícito, Tortura, Lesiones personales, Delitos contra la libertad sexual. Cada hecho con indicador de disponibilidad de ruta administrativa.')

# — Módulo 02
h2('Módulo 02 — Barreras desde el Territorio', '1C1F3A')
route_tag('/modulos/barreras-territorio')
body('Análisis de las barreras que enfrentan las víctimas para acceder a la indemnización, clasificadas en cuatro dimensiones.')
h3('Secciones del módulo:')
bullet('Cuatro tipos de barreras — Cards con color, ícono, descripción, sub-barreras e impacto estadístico:')
bullet('Geográficas — Distancias, vías, transporte, grupos armados, clima. Impacto: 62% en zonas rurales dispersas.', 1)
bullet('Administrativas — Tiempos, requisitos, cobertura, horarios. Impacto: 48% reportan al menos una barrera administrativa.', 1)
bullet('Sociales — Estigmatización, desconfianza, lengua, redes. Impacto: 71% de comunidades étnicas reportan barreras culturales.', 1)
bullet('De información — Desconocimiento de derechos, conectividad, alfabetismo. Impacto: 55% no conoce el monto de su indemnización.', 1)
bullet('Dashboard dinámico por departamento — Intensidad de barreras en escala 1–10 por departamento. Selección interactiva y comparación entre tipos de barrera.')

# — Módulo 03
h2('Módulo 03 — La Mirada desde las Víctimas', '1C1F3A')
route_tag('/modulos/mirada-victimas')
body('Perspectivas, experiencias y valoraciones de las víctimas sobre el proceso de reparación.')
h3('Secciones del módulo:')
bullet('Cuatro dimensiones de experiencia — Cards con datos cualitativos y cuantitativos:')
bullet('Percepción de barreras — Testimonios sobre obstáculos. Dato: 78% reportó al menos una barrera significativa.', 1)
bullet('Insatisfacción con la atención — Experiencias con calidad del servicio. Dato: 61% califica la atención como deficiente.', 1)
bullet('Desconfianza institucional — Niveles de confianza en el Estado. Testimonios sobre confiabilidad gubernamental.', 1)
bullet('Acceso a resultados — Seguimiento y accesibilidad de resultados del proceso.', 1)
bullet('Datos cualitativos — Citas textuales de víctimas integradas con métricas cuantitativas.')

# — Módulo 04
h2('Módulo 04 — Trayectorias de Contactabilidad y Localización', '1C1F3A')
route_tag('/modulos/trayectorias-contactabilidad')
body('Análisis operativo de la capacidad institucional para localizar y contactar a las víctimas en 15 departamentos de intervención CISP.')
h3('Secciones del módulo:')
bullet('Panel de departamentos — Grid con cards de los 15 departamentos: Antioquia, Atlántico, Bolívar, Boyacá, Cauca, Cesar, Chocó, Cundinamarca, Huila, Magdalena, Meta, Nariño, Risaralda, Santander, Valle del Cauca.')
bullet('Métricas por departamento — Tasa de éxito telefónico, % de teléfonos activos, intentos de contacto promedio, jornadas de atención, víctimas atendidas, municipios con cobertura, % de cobertura general.')
bullet('Panel de detalle — Al seleccionar un departamento se amplían las métricas con visualizaciones.')
bullet('Niveles de efectividad — Color-coded: Alto ≥70% (verde), Medio ≥50% (amarillo), Bajo ≥30% (naranja), Muy bajo <30% (rojo).')

# — Módulo 05
h2('Módulo 05 — Dinámicas Operativas Territoriales', '1C1F3A')
route_tag('/modulos/dinamicas-operativas')
body('Dashboard basado en el Instrumento 3 (levantamiento institucional). Visualiza las dinámicas operativas de las instituciones que atienden a víctimas en los territorios. 180 registros simulados de instituciones.')
h3('KPIs nacionales (fila superior):')
bullet('180 registros institucionales')
bullet('83% de casos con barreras identificadas')
bullet('52% de casos en avance')
bullet('35% con complejidad alta')
h3('Cuatro pestañas de análisis:')
bullet('Operación (F09–F13) — Distribución de personas atendidas por institución, niveles de complejidad del proceso, % que realiza acciones de indemnización, % que hace seguimientos.')
bullet('Barreras (F14–F25) — 8 tipos de barreras operativas con frecuencia e impacto: Contactabilidad (74%), Información incompleta (68%), Corrección de documentos (62%), Disponibilidad institucional (49%), Coordinación interinstitucional (44%), Conectividad (38%), Seguridad (31%), Acceso territorial (29%). Distribución de frecuencia (Muy frecuente/Frecuente/Ocasional/Aislado) e impacto (Alto/Medio/Bajo).')
bullet('Contacto y Territorio (F26–F39) — Análisis de canales de contacto, cobertura territorial y modalidades de atención.')
bullet('Resultados (F40–F47) — Análisis de resultados y avances en los procesos de indemnización.')
h3('Panel de ranking departamental:')
bullet('Lista de 15 departamentos ordenados por % de barreras.')
bullet('Al seleccionar un departamento se muestra su perfil detallado en el panel de la derecha.')
bullet('Panel de resumen nacional cuando no hay departamento seleccionado.')

# — ComprendiendoReparacion / otros
h2('5. Página No Encontrada (404)', '1C1F3A')
route_tag('* (catch-all)')
body('Página de error con mensaje "Página no encontrada" y botón de regreso al inicio.')

# ─── Pie del documento ──────────────────────────────────────────────────────
page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('CISP Colombia — Estrategia de Reparación Individual')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x9D, 0x9D, 0x9C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documento generado automáticamente a partir del código fuente de los proyectos.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ─── Guardar ─────────────────────────────────────────────────────────────────
out = r'c:\Users\jcsal\OneDrive\Documentos\dev\observatorio_cisp_v2\Documentacion_Sitios_CISP.docx'
doc.save(out)
print(f'Documento guardado: {out}')
